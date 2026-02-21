"""
Tool implementations for the Revenue Pipeline Engine.

Five tools available to every revenue stream agent:
  search_web     – Tavily-powered web search
  read_webpage   – Full text extraction via Jina Reader
  create_file    – Write txt/md/csv/json/docx/pdf to output dir
  execute_python – Run Python code (pandas, openpyxl, pptx, matplotlib…)
  finish         – Signal task complete + declare deliverables
"""

import os
import json
import subprocess
import tempfile
from typing import Any, Dict, List


# ---------------------------------------------------------------------------
# OpenAI function-calling tool definitions
# ---------------------------------------------------------------------------

TOOL_DEFINITIONS: List[Dict] = [
    {
        "type": "function",
        "function": {
            "name": "search_web",
            "description": (
                "Search the internet for up-to-date information. "
                "Returns a list of results with title, URL, and a content snippet."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "query": {
                        "type": "string",
                        "description": "Search query",
                    },
                    "max_results": {
                        "type": "integer",
                        "description": "Number of results (default 5, max 10)",
                        "default": 5,
                    },
                },
                "required": ["query"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "read_webpage",
            "description": (
                "Read the full text content of a webpage. "
                "Returns clean markdown-like text, ideal for research and summarization."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "url": {
                        "type": "string",
                        "description": "Full URL to fetch",
                    },
                },
                "required": ["url"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "create_file",
            "description": (
                "Create a file in the session output directory. "
                "Supported types: txt, md, csv, json, docx, pdf. "
                "Returns the absolute file path."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "filename": {
                        "type": "string",
                        "description": "Filename without extension",
                    },
                    "content": {
                        "type": "string",
                        "description": "Full file content (text or markdown)",
                    },
                    "file_type": {
                        "type": "string",
                        "enum": ["txt", "md", "csv", "json", "docx", "pdf"],
                        "description": "Output format",
                    },
                },
                "required": ["filename", "content", "file_type"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "execute_python",
            "description": (
                "Execute Python code to generate complex artifacts. "
                "Available libraries: pandas, openpyxl, xlsxwriter, python-pptx, "
                "matplotlib, reportlab, python-docx, Pillow, requests. "
                "The variable `output_dir` is pre-set to the session directory. "
                "Print 'FILE:<path>' for each file you create so the engine tracks them."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "code": {
                        "type": "string",
                        "description": (
                            "Python code to execute. Save output files inside `output_dir`. "
                            "Print FILE:<path> for each file created."
                        ),
                    },
                },
                "required": ["code"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "finish",
            "description": (
                "Signal that all deliverables are complete. "
                "Call this as the LAST step once all files have been created."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "summary": {
                        "type": "string",
                        "description": "1-3 sentence summary of what was accomplished",
                    },
                    "deliverables": {
                        "type": "array",
                        "items": {"type": "string"},
                        "description": "List of absolute file paths created during this session",
                    },
                },
                "required": ["summary", "deliverables"],
            },
        },
    },
]


# ---------------------------------------------------------------------------
# Tool dispatcher
# ---------------------------------------------------------------------------

def execute_tool(tool_name: str, tool_args: Dict[str, Any], output_dir: str) -> Any:
    """Route a tool call to the appropriate implementation."""
    dispatch = {
        "search_web": _search_web,
        "read_webpage": _read_webpage,
        "create_file": _create_file,
        "execute_python": _execute_python,
    }
    handler = dispatch.get(tool_name)
    if handler is None:
        return {"error": f"Unknown tool: {tool_name}"}

    if tool_name == "create_file":
        return handler(
            tool_args.get("filename", "output"),
            tool_args.get("content", ""),
            tool_args.get("file_type", "txt"),
            output_dir,
        )
    elif tool_name == "execute_python":
        return handler(tool_args.get("code", ""), output_dir)
    elif tool_name == "search_web":
        return handler(tool_args.get("query", ""), tool_args.get("max_results", 5))
    elif tool_name == "read_webpage":
        return handler(tool_args.get("url", ""))
    else:
        return handler(**tool_args)


# ---------------------------------------------------------------------------
# Individual tool implementations
# ---------------------------------------------------------------------------

def _search_web(query: str, max_results: int = 5) -> Dict:
    """Web search via Tavily API with fallback to DuckDuckGo."""
    api_key = os.getenv("WEB_SEARCH_API_KEY") or os.getenv("TAVILY_API_KEY")

    if api_key:
        try:
            from tavily import TavilyClient
            client = TavilyClient(api_key=api_key)
            resp = client.search(query=query, max_results=min(max_results, 10))
            results = [
                {
                    "title": r.get("title", ""),
                    "url": r.get("url", ""),
                    "snippet": r.get("content", "")[:600],
                }
                for r in resp.get("results", [])
            ]
            return {"query": query, "results": results, "source": "tavily"}
        except Exception as exc:
            return {"query": query, "results": [], "error": str(exc)}

    # Fallback: DuckDuckGo instant answer API (no key required)
    try:
        import requests
        resp = requests.get(
            "https://api.duckduckgo.com/",
            params={"q": query, "format": "json", "no_html": 1},
            timeout=10,
        )
        data = resp.json()
        results = []
        for topic in data.get("RelatedTopics", [])[:max_results]:
            if isinstance(topic, dict) and "Text" in topic:
                results.append({
                    "title": topic.get("Text", "")[:100],
                    "url": topic.get("FirstURL", ""),
                    "snippet": topic.get("Text", ""),
                })
        return {"query": query, "results": results, "source": "duckduckgo"}
    except Exception as exc:
        return {"query": query, "results": [], "error": f"No search provider available: {exc}"}


def _read_webpage(url: str) -> Dict:
    """
    Fetch clean text from a URL.
    Tries Jina Reader first (best quality), falls back to raw requests + BeautifulSoup.
    """
    try:
        import requests

        jina_key = os.getenv("WEB_SEARCH_API_KEY") or os.getenv("JINA_API_KEY")
        headers = {"Accept": "text/plain", "X-Return-Format": "markdown"}
        if jina_key and os.getenv("WEB_SEARCH_PROVIDER", "tavily") == "jina":
            headers["Authorization"] = f"Bearer {jina_key}"

        jina_url = f"https://r.jina.ai/{url}"
        resp = requests.get(jina_url, headers=headers, timeout=30)
        if resp.status_code == 200 and len(resp.text) > 100:
            return {"url": url, "content": resp.text[:10000], "source": "jina"}

        # Fallback: direct fetch + BeautifulSoup
        resp = requests.get(
            url,
            timeout=15,
            headers={"User-Agent": "Mozilla/5.0 (compatible; RevenueBot/1.0)"},
        )
        try:
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(resp.text, "html.parser")
            for tag in soup(["script", "style", "nav", "footer", "header", "aside"]):
                tag.decompose()
            content = soup.get_text(separator="\n", strip=True)
        except ImportError:
            content = resp.text

        return {"url": url, "content": content[:10000], "source": "requests"}

    except Exception as exc:
        return {"url": url, "content": "", "error": str(exc)}


def _create_file(filename: str, content: str, file_type: str, output_dir: str) -> Dict:
    """
    Write content to a file in output_dir.
    Handles plain text, markdown, CSV, JSON, DOCX, and PDF.
    """
    os.makedirs(output_dir, exist_ok=True)

    # Sanitize filename
    safe_name = "".join(c for c in filename if c.isalnum() or c in " ._-").strip()
    safe_name = safe_name or "output"
    file_path = os.path.join(output_dir, f"{safe_name}.{file_type}")

    try:
        if file_type in ("txt", "md", "csv", "json"):
            with open(file_path, "w", encoding="utf-8") as f:
                f.write(content)

        elif file_type == "docx":
            from docx import Document
            from docx.shared import Pt, RGBColor
            doc = Document()
            for line in content.split("\n"):
                stripped = line.strip()
                if stripped.startswith("### "):
                    doc.add_heading(stripped[4:], level=3)
                elif stripped.startswith("## "):
                    doc.add_heading(stripped[3:], level=2)
                elif stripped.startswith("# "):
                    doc.add_heading(stripped[2:], level=1)
                elif stripped.startswith("- ") or stripped.startswith("* "):
                    p = doc.add_paragraph(stripped[2:], style="List Bullet")
                elif stripped == "":
                    doc.add_paragraph("")
                else:
                    doc.add_paragraph(stripped)
            doc.save(file_path)

        elif file_type == "pdf":
            try:
                from reportlab.lib.pagesizes import letter
                from reportlab.platypus import (
                    SimpleDocTemplate, Paragraph, Spacer, HRFlowable
                )
                from reportlab.lib.styles import getSampleStyleSheet
                from reportlab.lib.units import inch

                doc = SimpleDocTemplate(file_path, pagesize=letter,
                                        leftMargin=inch, rightMargin=inch,
                                        topMargin=inch, bottomMargin=inch)
                styles = getSampleStyleSheet()
                story = []

                for line in content.split("\n"):
                    stripped = line.strip()
                    if stripped.startswith("# "):
                        story.append(Paragraph(stripped[2:], styles["h1"]))
                        story.append(HRFlowable(width="100%", thickness=1))
                    elif stripped.startswith("## "):
                        story.append(Paragraph(stripped[3:], styles["h2"]))
                    elif stripped.startswith("### "):
                        story.append(Paragraph(stripped[4:], styles["h3"]))
                    elif stripped:
                        story.append(Paragraph(stripped, styles["Normal"]))
                    story.append(Spacer(1, 0.08 * inch))

                doc.build(story)

            except Exception as pdf_exc:
                # Graceful fallback: save as .txt
                fallback_path = file_path.replace(".pdf", "_fallback.txt")
                with open(fallback_path, "w", encoding="utf-8") as f:
                    f.write(content)
                return {
                    "file_path": fallback_path,
                    "status": "created_as_txt",
                    "note": f"PDF generation failed ({pdf_exc}), saved as txt",
                    "size_bytes": os.path.getsize(fallback_path),
                }

        return {
            "file_path": file_path,
            "status": "created",
            "type": file_type,
            "size_bytes": os.path.getsize(file_path),
        }

    except Exception as exc:
        return {"file_path": None, "error": str(exc)}


def _execute_python(code: str, output_dir: str) -> Dict:
    """
    Execute Python code in a subprocess.
    Injects `output_dir` into the global scope.
    Parses FILE:<path> markers from stdout to track created files.
    """
    os.makedirs(output_dir, exist_ok=True)

    wrapper = f"""import os, sys
output_dir = {repr(output_dir)}
os.makedirs(output_dir, exist_ok=True)

{code}
"""
    with tempfile.NamedTemporaryFile(suffix=".py", mode="w", delete=False, encoding="utf-8") as tmp:
        tmp.write(wrapper)
        tmp_path = tmp.name

    try:
        result = subprocess.run(
            ["python", tmp_path],
            capture_output=True,
            text=True,
            timeout=90,
        )
        stdout = result.stdout or ""
        stderr = result.stderr or ""

        # Extract file paths from FILE: markers
        file_paths = []
        for line in stdout.splitlines():
            line = line.strip()
            if line.startswith("FILE:"):
                path = line[5:].strip()
                if os.path.exists(path):
                    file_paths.append(path)
            elif line.startswith("ARTIFACT_PATH:"):
                path = line[14:].strip()
                if os.path.exists(path):
                    file_paths.append(path)

        return {
            "stdout": stdout[:4000],
            "stderr": stderr[:1000],
            "exit_code": result.returncode,
            "file_paths": file_paths,
            "file_path": file_paths[0] if file_paths else None,
        }

    except subprocess.TimeoutExpired:
        return {"error": "Code execution timed out (90s)", "file_paths": [], "stdout": "", "stderr": ""}
    except Exception as exc:
        return {"error": str(exc), "file_paths": [], "stdout": "", "stderr": ""}
    finally:
        try:
            os.unlink(tmp_path)
        except OSError:
            pass
